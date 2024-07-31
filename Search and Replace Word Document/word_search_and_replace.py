import json

import docx
from docx.table import Table
from docx.shared import Mm
from typing import TYPE_CHECKING, List, Dict, Optional

if TYPE_CHECKING:
    from docx.document import Document as DocumentObject
    from docx.text.paragraph import Paragraph 
    from docx.section import _Header

def replace_text_in_paragraph(paragraph:"Paragraph", search_text, replace_text):
  if search_text in paragraph.text:
    runs = paragraph.runs
    for run in runs:
      if search_text in run.text:
        run.text = run.text.replace(search_text, replace_text)
    while paragraph.text.count(search_text) > 0:
      runs = paragraph.runs
      start = 0
      end = len(runs)
      while search_text in ''.join(run.text for run in runs[start+1:end]):
        start += 1
      while search_text in ''.join(run.text for run in runs[start:end-1]):
        end -= 1
      text = ''.join(run.text for run in runs[start:end])
      before_text, after_text = text.split(search_text, 1)
      for run in runs[start+1:end]:
        paragraph._p.remove(run._r)
      paragraph.runs[start].text = before_text
      run = paragraph.add_run(replace_text)
      paragraph.runs[start]._r.addnext(run._r)
      run = paragraph.add_run(after_text)
      paragraph.runs[start+1]._r.addnext(run._r)

def replace_text_in_table(table:Table, search_text, replace_text):
  for row in table.rows:
    for cell in row.cells:
      for paragraph in cell.paragraphs:
        replace_text_in_paragraph(paragraph, search_text, replace_text)

def replace_text_in_header(header:"_Header", search_text, replace_text):
  for paragraph in header.paragraphs:
    replace_text_in_paragraph(paragraph, search_text, replace_text)
  for table in header.tables:
    replace_text_in_table(table, search_text, replace_text)

def replace_text_in_document(doc:"DocumentObject", search_text, replace_text):
  for paragraph in doc.paragraphs:
    replace_text_in_paragraph(paragraph, search_text, replace_text)
  for table in doc.tables:
    replace_text_in_table(table, search_text, replace_text)
  
  for section in doc.sections:
    header = section.header
    replace_text_in_header(header, search_text, replace_text)

def add_image_in_paragraph(paragraph:"Paragraph", search_text, image_path, image_size):
  while paragraph.text.count(search_text) > 0:
    runs = paragraph.runs
    start = 0
    end = len(runs)
    while search_text in ''.join(run.text for run in runs[start+1:end]):
      start += 1
    while search_text in ''.join(run.text for run in runs[start:end-1]):
      end -= 1
    text = ''.join(run.text for run in runs[start:end])
    before_text, after_text = text.split(search_text, 1)
    for run in runs[start+1:end]:
      paragraph._p.remove(run._r)

    paragraph.runs[start].text = before_text
    run = paragraph.add_run()
    run.add_picture(image_path, Mm(image_size))
    paragraph.runs[start]._r.addnext(run._r)
    if after_text:
      run = paragraph.add_run(after_text)
      paragraph.runs[start+1]._r.addnext(run._r)
    if not before_text:
      paragraph._p.remove(paragraph.runs[start]._r)

def add_image_in_table(table:Table, search_text, image_path, image_size):
  for row in table.rows:
    for cell in row.cells:
      for paragraph in cell.paragraphs:
        add_image_in_paragraph(paragraph, search_text, image_path, image_size)

def add_image_in_header(header:"_Header", search_text, image_path, image_size):
  for paragraph in header.paragraphs:
    add_image_in_paragraph(paragraph, search_text, image_path, image_size)
  for table in header.tables:
    add_image_in_table(table, search_text, image_path, image_size)

def add_image_in_document(doc:"DocumentObject", search_text, image_path, image_size):
  for paragraph in doc.paragraphs:
    add_image_in_paragraph(paragraph, search_text, image_path, image_size)
  for table in doc.tables:
    add_image_in_table(table, search_text, image_path, image_size)

  for section in doc.sections:
    header = section.header
    add_image_in_header(header, search_text, image_path, image_size)

def replace_text_and_add_images(doc:"DocumentObject", replacements: List[Dict[str, Optional[str]]]):
  replace_texts = replacements["text"]
  for replace_text in replace_texts:
    search_text = replace_text.get('search_text')
    replace_text = replace_text.get('replace_text')
    replace_text_in_document(doc, search_text, replace_text)

  replace_images = replacements["image"]
  for replace_image in replace_images: 
    search_text = replace_image.get('search_text')
    image_path = replace_image.get('image_path')
    image_size = float(replace_image.get('image_size_in_mm'))
    add_image_in_document(doc, search_text, image_path, image_size)
        


#Main part of the code

#Open template file
document = docx.Document("Word template.docx")

#Open and read json file. Can also get replacements from other sources like Excel
with open('replacement.json', 'r') as file:
  replacements = json.load(file)


replace_text_and_add_images(document, replacements)

#Save the new word document
document.save("Output.docx")

         
